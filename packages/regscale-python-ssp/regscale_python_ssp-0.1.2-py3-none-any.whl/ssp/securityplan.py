import re

from docx import Document
from ssp.control import Control
from lxml import etree
from io import StringIO

class SecurityPlan(object):
    """
    Not intended to be constructed directly. Use ssp.api.SSP() to open an SSP.
    """

    def __init__(self, path):
        self.source = path
        self.control_list = []
        self.control_list_to_table_index = {}
        self.system_security_level = None
        self.document = Document(path)
        self.version = self.get_version()
        self.child_ssp = None
        if self.version == '08282018':
            self.child_ssp = SecurityPlan_08282018(self)
        else:
            raise ValueError('This template version is not compatible.')
        self.revision = self.child_ssp.revision()
        self.control_list = self.child_ssp.control_list
        self.control_list_to_table_index = self.child_ssp.control_list_to_table_index
        self.system_security_level = self.child_ssp.system_security_level()
        self.security_categorization = self.child_ssp.security_categorization()

    def __iter__(self):
        return iter(self.control_list)

    def control(self, control):
        """
        Takes a control as a string and returns the matching control object in the SSP.
        """
        try:
            return self.control_list[self.control_list_to_table_index[control.upper()]]
        except KeyError:
            raise KeyError('No control found with that name')

    def get_version(self):
        """
        Needs to take a Document file and return either a version number or raise an error if not an ssp.
        """
        # For some reason docx doesn't see anything before the table of contents in the list of tables or paragraphs, will need to parse document._element.xml
        # template_revision_table_index = self.table_name_to_table_index['Date']
        # template_revision_table = self.tables[template_revision_table_index]
        # try:
        #     date = template_revision_table.cell(7, 0).text
        #     if date == '8/28/2018':
        #         return '08282018'
        #     else:
        #         raise ValueError
        # except:
        #     raise ValueError('Error getting version, has the template revision history been altered? Possibly unsupported template')
        return '08282018'

class SecurityPlan_08282018(SecurityPlan):
    """
    Child class for the SSP version published 08/28/2018. 
    Takes an SSP object and returns a child SSP object specific to this template.
    """

    def __init__(self, ssp):
        self.control_list = []
        self.control_list_to_table_index = {}
        self.document = ssp.document
        self.tables = ssp.document.tables
        self.create_control_index()
        self.namespaces = {
            "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
            "w14": "http://schemas.microsoft.com/office/word/2010/wordml",
        }
    def is_cis_table(self, table):
        try:
            if '-' in table.cell(0,0).text and self.is_valid_control(table.cell(0,0).text):
                return True
            return False
        except Exception:
            return False

    def is_valid_control(self, text):
        """
        Takes control text (cis_table cell 0,0) and returns whether it is properly formatted.
        Regex needs to match AC-1 / AC-2(2) / AC-2(Ext) / AC-2(Privacy)
        etc all with/without spaces.
        TODO: In the future could use this to collect more properties like:
        - is it a privacy control or extension?
        - is it an enhancement? 
        """
        regex = r'([A-Z]*-[0-9]*\s*(\([0-9]*\))*)(\s*\([A-z]*\))*\s*(Req.)*$'
        match_groups = re.search(regex, text)
        return bool(match_groups)

    def create_control_index(self):
        """
        Creates control objects for each pair of CIS/Implementation tables
        and adds them to control_list and the list index.
        """
        implementation_table_next = False
        cis_table = ''
        for table in self.document.tables:
            try:
                if self.is_cis_table(table):
                    implementation_table_next = True
                    cis_table = table
                elif implementation_table_next:
                    new_control = Control(cis_table, table)
                    self.control_list_to_table_index[new_control.number.upper()] = len(self.control_list)
                    self.control_list.append(new_control)
                    implementation_table_next = False
            except Exception as e:
                print(e)

    def system_security_level(self) -> str:
        """The System Security Level
        """    
        result = None   
        dat =  self.get_xpath_data(key='System Sensitivity Level')
        if dat:
            result = dat[0]
        return result

    def security_categorization(self) -> str:
        """The Security Categorization.
        """     
        result = None   
        dat = self.get_xpath_data(key='Security Categorization')
        if dat:
            result = dat[0]
        return result

    def revision(self) -> str:
        """The SSP revision.
        """
        result = None
        regex = r'Version\s#*([\d.]+)'  
        try:
            fed_ramp_revision_string = (self.document.tables[0].cell(0,0).text) # First cell of the first table.
            result =  re.search(regex, fed_ramp_revision_string).group(1)
        except (AttributeError, IndexError):
            print('Warning, unable to return revision information.')
        return result
            
    def get_xpath_data(self, key, xpath='//w:tbl/w:tr/w:sdt/w:sdtContent/w:tc/w:p/w:r/w:t'):
        """_summary_

        :param key: _description_
        :param xpath: _description_
        :return: _description_
        """        
        tables = iter(self.document.tables)
        for _, t in enumerate(tables):
            if key in t._element.xml:
                f = StringIO(t._element.xml)
                tree = etree.parse(f)
                r = tree.xpath(xpath, namespaces=self.namespaces)
                if r:
                    txt = r[0].text.split()  # Remove newlines and spaces
                    if txt:
                        return txt
                return None